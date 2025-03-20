#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Módulo para geração avançada de documentos DOCX
"""

import os
import re
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION, WD_SECTION

# Importar funções de formatação jurídica
from .formatacao_juridica import (
    configurar_estilos_juridicos,
    formatar_citacoes_legais,
    adicionar_numeracao_paginas,
    formatar_jurisprudencia,
    adicionar_marca_dagua,
    formatar_texto_juridico
)

class DocxGenerator:
    """Classe para geração avançada de documentos DOCX"""
    
    def __init__(self, templates_dir, peticoes_dir, clientes_dir):
        """Inicializa o gerador de documentos DOCX"""
        self.templates_dir = templates_dir
        self.peticoes_dir = peticoes_dir
        self.clientes_dir = clientes_dir
        
        # Criar diretórios se não existirem
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.peticoes_dir, exist_ok=True)
        os.makedirs(self.clientes_dir, exist_ok=True)
        
        # Mapeamento de tipos de petição para nomes de templates
        self.tipo_template_map = {
            "recurso": "recurso_administrativo",
            "recurso administrativo": "recurso_administrativo",
            "impugnação": "impugnacao_edital",
            "impugnação ao edital": "impugnacao_edital",
            "mandado": "mandado_seguranca",
            "mandado de segurança": "mandado_seguranca",
            "contrarrazões": "contrarrazoes_recurso",
            "contrarrazões de recurso": "contrarrazoes_recurso"
        }
    
    def _obter_template_path(self, tipo):
        """Obtém o caminho do template com base no tipo de petição"""
        tipo_lower = tipo.lower()
        tipo_template = self.tipo_template_map.get(tipo_lower, tipo_lower.replace(' ', '_'))
        template_path = os.path.join(self.templates_dir, f"{tipo_template}.docx")
        
        # Verificar se o template existe
        if not os.path.exists(template_path):
            # Tentar encontrar um template alternativo
            templates = [f for f in os.listdir(self.templates_dir) 
                         if f.endswith('.docx') and not f.startswith('~$')]
            
            # Tentar encontrar um template que contenha parte do nome do tipo
            matching_templates = [t for t in templates if tipo_lower in t.lower()]
            if matching_templates:
                template_path = os.path.join(self.templates_dir, matching_templates[0])
            elif templates:
                template_path = os.path.join(self.templates_dir, templates[0])
            else:
                template_path = None
        
        return template_path
    
    def _obter_dados_cliente(self, cliente_id):
        """Obtém os dados do cliente pelo ID"""
        if not cliente_id:
            return None
        
        print(f"Buscando dados do cliente com ID: {cliente_id}")
        
        # Carregar informações do cliente
        clientes_json_path = os.path.join(self.clientes_dir, 'clientes.json')
        if os.path.exists(clientes_json_path):
            try:
                with open(clientes_json_path, 'r', encoding='utf-8') as f:
                    clientes = json.load(f)
                    # Converter cliente_id para string para comparação
                    cliente_id_str = str(cliente_id)
                    cliente = next((c for c in clientes if str(c.get('id')) == cliente_id_str), None)
                    
                    if cliente:
                        print(f"Cliente encontrado: {cliente.get('nome', 'Nome não disponível')}")
                    else:
                        print(f"Cliente com ID {cliente_id} não encontrado no arquivo clientes.json")
                    
                    return cliente
            except Exception as e:
                print(f"Erro ao carregar dados do cliente: {str(e)}")
        else:
            print(f"Arquivo de clientes não encontrado: {clientes_json_path}")
        
        return None
    
    def _obter_logo_cliente(self, cliente_nome):
        """Obtém o caminho da logo do cliente pelo nome"""
        if not cliente_nome:
            return None
        
        # Buscar logo pelo nome do cliente
        logos_dir = os.path.join(self.clientes_dir, 'logos')
        if os.path.exists(logos_dir):
            # Procurar por arquivos de imagem com o nome do cliente
            for ext in ['.png', '.jpg', '.jpeg', '.gif']:
                potential_logo = os.path.join(
                    logos_dir, 
                    f"{cliente_nome.lower().replace(' ', '_')}{ext}"
                )
                if os.path.exists(potential_logo):
                    return potential_logo
        
        return None
    
    def _processar_texto_juridico(self, texto):
        """Processa o texto para formatação jurídica"""
        if not texto:
            return ""
        
        # Aplicar formatações jurídicas
        texto = formatar_texto_juridico(texto)
        
        # Remover espaços em branco extras
        texto = re.sub(r'\s+', ' ', texto).strip()
        
        return texto
    
    def _substituir_placeholders(self, doc, dados):
        """Substitui placeholders no documento"""
        # Extrair dados
        fatos = self._processar_texto_juridico(dados.get('fatos', ''))
        fundamentos = self._processar_texto_juridico(dados.get('fundamentos', ''))
        pedidos = self._processar_texto_juridico(dados.get('pedidos', ''))
        cliente_nome = dados.get('cliente_nome', 'Cliente')
        cliente_qualificacao = dados.get('cliente_qualificacao', '')
        advogado_nome = dados.get('advogado_nome', 'ADVOGADO')
        advogado_oab = dados.get('advogado_oab', 'OAB/XX 12345')
        cidade = dados.get('cidade', 'São Paulo')
        autoridade = dados.get('autoridade', '')
        referencia_processo = dados.get('referencia_processo', '')
        logo_path = dados.get('logo_path')
        
        # Verificar se os dados estão corretos
        print(f"Substituindo placeholders com os seguintes dados:")
        print(f"  Cliente: {cliente_nome}")
        print(f"  Qualificação: {cliente_qualificacao}")
        print(f"  Autoridade: {autoridade}")
        print(f"  Fatos: {bool(fatos)}")
        print(f"  Fundamentos: {bool(fundamentos)}")
        print(f"  Pedidos: {bool(pedidos)}")
        
        # Verificar se é um Recurso Administrativo
        is_recurso = False
        for paragraph in doc.paragraphs:
            if "RECURSO ADMINISTRATIVO" in paragraph.text:
                is_recurso = True
                break
        
        # Mapeamento de placeholders para valores
        placeholders = {
            "[FATOS]": fatos,
            "##FATOS##": fatos,
            "[FUNDAMENTOS]": fundamentos,
            "##FUNDAMENTOS##": fundamentos,
            "##ARGUMENTOS##": fundamentos,
            "[PEDIDOS]": pedidos,
            "##PEDIDOS##": pedidos,
            "##PEDIDO##": pedidos,
            "[DATA]": datetime.now().strftime('%d/%m/%Y'),
            "##DATA##": datetime.now().strftime('%d/%m/%Y'),
            "[CIDADE]": cidade,
            "##CIDADE##": cidade,
            "[AUTORIDADE]": autoridade,
            "##AUTORIDADE##": autoridade,
            "[REFERENCIA_PROCESSO]": referencia_processo,
            "##REFERENCIA_PROCESSO##": referencia_processo,
            "##PROCESSO##": referencia_processo,
            "[NOME_CLIENTE]": cliente_nome,
            "##NOME_CLIENTE##": cliente_nome,
            "##NOME_RECORRENTE##": cliente_nome,
            "[QUALIFICACAO_CLIENTE]": cliente_qualificacao,
            "##QUALIFICACAO_CLIENTE##": cliente_qualificacao,
            "[ADVOGADO]": advogado_nome,
            "##ADVOGADO##": advogado_nome,
            "##NOME_ADVOGADO##": advogado_nome,
            "[NUMERO_OAB]": advogado_oab,
            "##NUMERO_OAB##": advogado_oab
        }
        
        # Substituir placeholders em parágrafos
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            
            # Se for Recurso Administrativo e contém "em face de [CONTRAPARTE]"
            if is_recurso and "em face de [CONTRAPARTE]" in original_text:
                # Substituir o parágrafo inteiro
                paragraph.text = "pelos fatos e fundamentos a seguir expostos."
                continue
            
            # Remover qualquer referência à contraparte no documento
            if "[CONTRAPARTE]" in original_text or "##CONTRAPARTE##" in original_text:
                paragraph.text = paragraph.text.replace("[CONTRAPARTE]", "").replace("##CONTRAPARTE##", "")
            
            # Substituir cada placeholder
            for placeholder, valor in placeholders.items():
                if placeholder in original_text:
                    # Se o placeholder for para fatos, fundamentos ou pedidos,
                    # criar um novo parágrafo para cada linha
                    if placeholder in ["[FATOS]", "##FATOS##", "[FUNDAMENTOS]", "##FUNDAMENTOS##", 
                                    "##ARGUMENTOS##", "[PEDIDOS]", "##PEDIDOS##", "##PEDIDO##"]:
                        # Limpar o parágrafo atual
                        paragraph.text = ""
                        
                        # Dividir o texto em linhas
                        linhas = valor.split('\n')
                        
                        # Adicionar cada linha como um novo run
                        for i, linha in enumerate(linhas):
                            if linha.strip():  # Se a linha não estiver vazia
                                if i > 0:  # Se não for a primeira linha, adicionar quebra de linha
                                    paragraph.add_run('\n')
                                run = paragraph.add_run(linha.strip())
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                    else:
                        # Para outros placeholders, fazer substituição simples
                        paragraph.text = paragraph.text.replace(placeholder, valor or '')
            
            # Tratar logo separadamente
            if "[LOGO_CLIENTE]" in original_text and logo_path:
                paragraph.text = ""
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Cm(5))
            elif "##LOGO_CLIENTE##" in original_text and logo_path:
                paragraph.text = ""
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Cm(5))
        
        # Substituir placeholders em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        
                        # Substituir cada placeholder
                        for placeholder, valor in placeholders.items():
                            if placeholder in original_text:
                                if placeholder in ["[FATOS]", "##FATOS##", "[FUNDAMENTOS]", "##FUNDAMENTOS##", 
                                                "##ARGUMENTOS##", "[PEDIDOS]", "##PEDIDOS##", "##PEDIDO##"]:
                                    # Limpar o parágrafo atual
                                    paragraph.text = ""
                                    
                                    # Dividir o texto em linhas
                                    linhas = valor.split('\n')
                                    
                                    # Adicionar cada linha como um novo run
                                    for i, linha in enumerate(linhas):
                                        if linha.strip():  # Se a linha não estiver vazia
                                            if i > 0:  # Se não for a primeira linha, adicionar quebra de linha
                                                paragraph.add_run('\n')
                                            run = paragraph.add_run(linha.strip())
                                            run.font.name = 'Arial'
                                            run.font.size = Pt(12)
                                else:
                                    # Para outros placeholders, fazer substituição simples
                                    paragraph.text = paragraph.text.replace(placeholder, valor or '')
        
        return doc
    
    def gerar_documento(self, tipo, dados_peticao):
        """
        Gera um documento DOCX com base no tipo e nos dados da petição
        
        Args:
            tipo: Tipo da petição (ex: 'recurso administrativo')
            dados_peticao: Dicionário com os dados da petição
                - fatos: Texto dos fatos
                - fundamentos: Texto dos fundamentos
                - pedidos: Texto dos pedidos
                - cliente_id: ID do cliente (opcional)
                - contraparte: Nome da contraparte (opcional)
                - autoridade: Nome da autoridade (opcional)
                - referencia_processo: Referência do processo (opcional)
                - cidade: Cidade para o fechamento (opcional)
                
        Returns:
            Caminho do arquivo gerado
        """
        try:
            print(f"Iniciando geração do documento DOCX para petição: {tipo}")
            
            # Verificar se dados_peticao é um dicionário válido
            if not dados_peticao or not isinstance(dados_peticao, dict):
                print(f"AVISO: dados_peticao inválido: {dados_peticao}")
                dados_peticao = {}
            
            # Obter o caminho do template
            template_path = self._obter_template_path(tipo)
            
            # Criar um novo documento a partir do template ou em branco
            if template_path and os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = Document()
                # Configurar estilos jurídicos padrão
                doc = configurar_estilos_juridicos(doc)
            
            # Adicionar numeração de páginas
            doc = adicionar_numeracao_paginas(doc)
            
            # Obter dados do cliente
            cliente_id = dados_peticao.get('cliente_id')
            cliente_nome = dados_peticao.get('cliente_nome')
            cliente_razao_social = dados_peticao.get('cliente_razao_social')
            cliente_cnpj = dados_peticao.get('cliente_cnpj')
            cliente = None
            try:
                cliente = self._obter_dados_cliente(cliente_id)
            except Exception as e:
                print(f"Erro ao obter dados do cliente: {e}")
            
            print(f"Dados da petição recebidos: {dados_peticao}")
            
            # Preparar dados para substituição com valores padrão seguros
            dados_substituicao = {
                'fatos': dados_peticao.get('fatos_texto', ''),
                'fundamentos': dados_peticao.get('argumentos_texto', ''),
                'pedidos': dados_peticao.get('pedidos_texto', ''),
                'cliente_nome': cliente_nome or 'Cliente',
                'cliente_qualificacao': '',
                'advogado_nome': 'ADVOGADO',
                'advogado_oab': 'OAB/XX 12345',
                'cidade': dados_peticao.get('cidade', 'São Paulo'),
                'contraparte': dados_peticao.get('contraparte', ''),
                'autoridade': dados_peticao.get('autoridade', ''),
                'referencia_processo': dados_peticao.get('referencia_processo', '')
            }
            
            # Atualizar com valores reais do cliente se disponíveis
            if cliente and isinstance(cliente, dict):
                dados_substituicao['cliente_nome'] = cliente_nome or cliente.get('nome', 'Cliente')
                
                # Adicionar qualificação do cliente
                cliente_cnpj = cliente_cnpj or cliente.get('cnpj', '')
                cliente_endereco = cliente.get('endereco', '')
                
                if cliente_cnpj and cliente_endereco:
                    dados_substituicao['cliente_qualificacao'] = (
                        f"pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {cliente_cnpj}, "
                        f"com sede na {cliente_endereco}"
                    )
                elif cliente_cnpj:
                    dados_substituicao['cliente_qualificacao'] = f"pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {cliente_cnpj}"
                elif cliente_endereco:
                    dados_substituicao['cliente_qualificacao'] = f"pessoa jurídica de direito privado, com sede na {cliente_endereco}"
                
                # Obter dados do advogado do cliente
                advogados = cliente.get('advogados', [])
                if advogados and len(advogados) > 0:
                    advogado = advogados[0]
                    dados_substituicao['advogado_nome'] = advogado.get('nome', 'ADVOGADO')
                    dados_substituicao['advogado_oab'] = advogado.get('oab', 'OAB/XX 12345')
            else:
                # Se não tiver cliente do banco de dados, usar os dados recebidos diretamente
                if cliente_nome:
                    dados_substituicao['cliente_nome'] = cliente_nome
                
                if cliente_cnpj:
                    dados_substituicao['cliente_qualificacao'] = f"pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {cliente_cnpj}"
            
            print(f"Dados para substituição: {dados_substituicao}")
            
            # Substituir valores específicos do formulário se fornecidos
            if dados_peticao.get('nome_advogado'):
                dados_substituicao['advogado_nome'] = dados_peticao.get('nome_advogado')
            if dados_peticao.get('numero_oab'):
                dados_substituicao['advogado_oab'] = dados_peticao.get('numero_oab')
            
            # Obter logo do cliente
            logo_path = None
            try:
                if cliente and isinstance(cliente, dict):
                    logo_path = cliente.get('logo_path')
                    if not logo_path:
                        logo_path = self._obter_logo_cliente(cliente.get('nome'))
            except Exception as e:
                print(f"Erro ao obter logo do cliente: {e}")
                
            dados_substituicao['logo_path'] = logo_path
            
            # Substituir placeholders
            try:
                doc = self._substituir_placeholders(doc, dados_substituicao)
            except Exception as e:
                print(f"Erro ao substituir placeholders: {e}")
            
            # Definir o nome do arquivo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            tipo_sanitizado = re.sub(r'[^a-zA-Z0-9]', '_', tipo)
            
            # Usar o nome do cliente no nome do arquivo se disponível
            cliente_nome_sanitizado = re.sub(r'[^a-zA-Z0-9]', '_', dados_substituicao['cliente_nome'])
            
            filename = f"{tipo_sanitizado}_{cliente_nome_sanitizado}_{timestamp}.docx"
            filepath = os.path.join(self.peticoes_dir, filename)
            
            # Salvar o documento
            doc.save(filepath)
            
            print(f"Documento DOCX gerado com sucesso: {filepath}")
            
            return filepath
            
        except Exception as e:
            print(f"Erro ao gerar documento DOCX: {e}")
            import traceback
            traceback.print_exc()
            raise e 