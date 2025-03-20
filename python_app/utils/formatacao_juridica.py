#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Módulo para formatação jurídica avançada de documentos
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def configurar_estilos_juridicos(doc):
    """Configura estilos jurídicos padrão no documento"""
    # Estilo para títulos
    if 'Título Jurídico' not in doc.styles:
        style = doc.styles.add_style('Título Jurídico', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Estilo para texto normal
    if 'Texto Jurídico' not in doc.styles:
        style = doc.styles.add_style('Texto Jurídico', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.first_line_indent = Cm(1.25)
    
    # Estilo para citações
    if 'Citação Jurídica' not in doc.styles:
        style = doc.styles.add_style('Citação Jurídica', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.left_indent = Cm(4)
        style.paragraph_format.right_indent = Cm(4)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Estilo para cabeçalho
    if 'Cabeçalho Jurídico' not in doc.styles:
        style = doc.styles.add_style('Cabeçalho Jurídico', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilo para rodapé
    if 'Rodapé Jurídico' not in doc.styles:
        style = doc.styles.add_style('Rodapé Jurídico', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def formatar_citacoes_legais(texto):
    """Formata citações legais no texto para um formato padronizado"""
    # Padronizar referências a artigos de lei
    texto = re.sub(r'(?i)art(?:igo)?\.?\s*(\d+)', r'Art. \1', texto)
    texto = re.sub(r'(?i)§\s*(\d+)', r'§ \1', texto)
    texto = re.sub(r'(?i)inciso\s*([ivxlcdm]+)', r'inciso \1', texto)
    
    # Padronizar referências a leis
    texto = re.sub(r'(?i)lei\s*(?:n[°º]?)?\s*(\d+[\.\d]*)/(\d{4})', r'Lei nº \1/\2', texto)
    texto = re.sub(r'(?i)decreto\s*(?:n[°º]?)?\s*(\d+[\.\d]*)/(\d{4})', r'Decreto nº \1/\2', texto)
    
    # Padronizar referências a códigos
    texto = re.sub(r'(?i)c[óo]digo\s+civil', r'Código Civil', texto)
    texto = re.sub(r'(?i)c[óo]digo\s+de\s+processo\s+civil', r'Código de Processo Civil', texto)
    texto = re.sub(r'(?i)c[óo]digo\s+penal', r'Código Penal', texto)
    texto = re.sub(r'(?i)c[óo]digo\s+tribut[áa]rio\s+nacional', r'Código Tributário Nacional', texto)
    
    # Padronizar referências a constituição
    texto = re.sub(r'(?i)constitui[çc][ãa]o\s+federal', r'Constituição Federal', texto)
    texto = re.sub(r'(?i)cf/88', r'Constituição Federal', texto)
    
    return texto

def adicionar_numeracao_paginas(doc):
    """Adiciona numeração de páginas ao documento"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._element.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    
    return doc

def formatar_jurisprudencia(texto_jurisprudencia, tribunal="STJ"):
    """Formata citações de jurisprudência no padrão ABNT"""
    # Exemplo: formatar_jurisprudencia("REsp 1234567/SP, Rel. Min. João Silva, Segunda Turma, julgado em 01/01/2020, DJe 15/01/2020")
    if not texto_jurisprudencia:
        return texto_jurisprudencia
    
    # Padronizar siglas de tribunais
    tribunais = {
        "STJ": "Superior Tribunal de Justiça",
        "STF": "Supremo Tribunal Federal",
        "TRF": "Tribunal Regional Federal",
        "TJ": "Tribunal de Justiça",
        "TST": "Tribunal Superior do Trabalho"
    }
    
    tribunal_completo = tribunais.get(tribunal, tribunal)
    
    # Formatar no padrão ABNT
    texto_formatado = f"{tribunal_completo}. {texto_jurisprudencia}"
    
    return texto_formatado

def adicionar_marca_dagua(doc, texto="RASCUNHO"):
    """Adiciona marca d'água ao documento"""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run(texto)
    run.font.size = Pt(72)
    run.font.color.rgb = RGBColor(200, 200, 200)
    
    return doc

def formatar_texto_juridico(texto):
    """Aplica formatações jurídicas padrão ao texto"""
    # Aplicar formatação de citações legais
    texto = formatar_citacoes_legais(texto)
    
    # Padronizar termos jurídicos comuns
    termos_juridicos = {
        r'(?i)a\s*quo': 'a quo',
        r'(?i)ad\s*quem': 'ad quem',
        r'(?i)habeas\s*corpus': 'Habeas Corpus',
        r'(?i)de\s*cujus': 'de cujus',
        r'(?i)ex\s*nunc': 'ex nunc',
        r'(?i)ex\s*tunc': 'ex tunc',
        r'(?i)in\s*limine': 'in limine',
        r'(?i)in\s*dubio\s*pro\s*reo': 'in dubio pro reo',
        r'(?i)data\s*venia': 'data venia',
        r'(?i)mutatis\s*mutandis': 'mutatis mutandis'
    }
    
    for padrao, substituicao in termos_juridicos.items():
        texto = re.sub(padrao, substituicao, texto)
    
    return texto 