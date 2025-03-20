import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Diretórios
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates', 'docx')
CLIENTES_DIR = os.path.join(BASE_DIR, 'data')

# Criar diretórios se não existirem
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(CLIENTES_DIR, exist_ok=True)

class Cliente:
    """Classe para gerenciar informações de clientes"""
    def __init__(self, id, nome, cnpj=None, endereco=None, logo_path=None, advogados=None, 
                 razao_social=None, representante_legal=None, endereco_completo=None):
        self.id = id
        self.nome = nome
        self.cnpj = cnpj
        self.endereco = endereco
        self.logo_path = logo_path
        self.advogados = advogados or []
        self.razao_social = razao_social
        self.representante_legal = representante_legal
        self.endereco_completo = endereco_completo
    
    @staticmethod
    def carregar_clientes():
        """Carrega todos os clientes do diretório de clientes"""
        clientes = []
        clientes_file = os.path.join(CLIENTES_DIR, 'clientes.json')
        
        if os.path.exists(clientes_file):
            try:
                with open(clientes_file, 'r', encoding='utf-8') as f:
                    clientes_data = json.load(f)
                    for cliente_data in clientes_data:
                        clientes.append(Cliente(
                            id=cliente_data.get('id'),
                            nome=cliente_data.get('nome'),
                            cnpj=cliente_data.get('cnpj'),
                            endereco=cliente_data.get('endereco'),
                            logo_path=cliente_data.get('logo_path'),
                            advogados=cliente_data.get('advogados', []),
                            razao_social=cliente_data.get('razao_social'),
                            representante_legal=cliente_data.get('representante_legal'),
                            endereco_completo=cliente_data.get('endereco_completo')
                        ))
            except Exception as e:
                print(f"Erro ao carregar clientes: {e}")
        
        return clientes
    
    @staticmethod
    def carregar_clientes_do_json():
        """Carrega clientes do arquivo JSON externo (src/data/clientes.json)"""
        clientes = []
        clientes_file = os.path.join(os.path.dirname(BASE_DIR), 'src', 'data', 'clientes.json')
        
        if os.path.exists(clientes_file):
            try:
                with open(clientes_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    clientes_data = data.get('results', [])
                    for cliente_data in clientes_data:
                        # Construir endereço completo
                        endereco_completo = f"{cliente_data.get('EnderecoRua', '')}, {cliente_data.get('EnderecoNumero', '')}, {cliente_data.get('EnderecoBairro', '')}, {cliente_data.get('EnderecoCidade', '')}/{cliente_data.get('EnderecoUF', '')}, CEP: {cliente_data.get('EnderecoCep', '')}"
                        
                        clientes.append(Cliente(
                            id=cliente_data.get('_id'),
                            nome=cliente_data.get('Nome_fantasia'),
                            cnpj=cliente_data.get('CNPJ'),
                            endereco=endereco_completo,
                            razao_social=cliente_data.get('Razao_social'),
                            endereco_completo=endereco_completo
                        ))
            except Exception as e:
                print(f"Erro ao carregar clientes do JSON externo: {e}")
        
        return clientes
    
    @staticmethod
    def obter_cliente_por_id(cliente_id):
        """Obtém um cliente pelo ID"""
        # Primeiro tenta carregar do arquivo interno
        clientes = Cliente.carregar_clientes()
        for cliente in clientes:
            if cliente.id == cliente_id:
                return cliente
        
        # Se não encontrar, tenta carregar do arquivo externo
        clientes = Cliente.carregar_clientes_do_json()
        for cliente in clientes:
            if cliente.id == cliente_id:
                return cliente
        
        return None

class TemplateManager:
    """Classe para gerenciar templates de petições"""
    
    TIPOS_PETICAO = {
        'recurso_administrativo': {
            'titulo': 'RECURSO ADMINISTRATIVO',
            'papel_cliente': 'RECORRENTE',
            'papel_contraparte': 'RECORRIDA'
        },
        'contrarrazoes_recurso': {
            'titulo': 'CONTRARRAZÕES AO RECURSO ADMINISTRATIVO',
            'papel_cliente': 'RECORRIDA',
            'papel_contraparte': 'RECORRENTE'
        },
        'mandado_seguranca': {
            'titulo': 'MANDADO DE SEGURANÇA',
            'papel_cliente': 'IMPETRANTE',
            'papel_contraparte': 'IMPETRADA'
        },
        'impugnacao_edital': {
            'titulo': 'IMPUGNAÇÃO AO EDITAL',
            'papel_cliente': 'IMPUGNANTE',
            'papel_contraparte': 'IMPUGNADA'
        }
    }
    
    @staticmethod
    def criar_template_base():
        """Cria um template base para todos os tipos de petição"""
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        
        # Adicionar espaço para cabeçalho (logo e informações do cliente)
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "[LOGO_CLIENTE]"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
    
    @staticmethod
    def criar_templates_padrao():
        """Cria templates padrão para os tipos de petição mais comuns"""
        for tipo_id, info in TemplateManager.TIPOS_PETICAO.items():
            doc = TemplateManager.criar_template_base()
            
            # Adicionar título
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"EXCELENTÍSSIMO(A) SENHOR(A) [AUTORIDADE]")
            run.bold = True
            run.font.size = Pt(12)
            
            # Adicionar referência ao processo
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[REFERENCIA_PROCESSO]")
            run.bold = True
            
            # Adicionar espaço
            doc.add_paragraph()
            
            # Adicionar qualificação
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run("[NOME_CLIENTE], [QUALIFICACAO_CLIENTE], vem, respeitosamente, à presença de Vossa Excelência, por intermédio de seu(sua) advogado(a) que esta subscreve, com fundamento em [BASE_LEGAL], interpor o presente")
            
            # Adicionar tipo de petição
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(info['titulo'])
            run.bold = True
            run.font.size = Pt(12)
            
            # Adicionar contraparte
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run("em face de [CONTRAPARTE], pelos fatos e fundamentos a seguir expostos.")
            
            # Adicionar seções
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("I - DOS FATOS")
            run.bold = True
            
            doc.add_paragraph("[FATOS]")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("II - DOS FUNDAMENTOS")
            run.bold = True
            
            doc.add_paragraph("[FUNDAMENTOS]")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("III - DOS PEDIDOS")
            run.bold = True
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run("Ante o exposto, requer:")
            
            doc.add_paragraph("[PEDIDOS]")
            
            # Adicionar fechamento
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run("Nestes termos,")
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run("Pede deferimento.")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run("[CIDADE], [DATA].")
            
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("[ADVOGADO]")
            p.add_run("\n")
            p.add_run("[NUMERO_OAB]")
            
            # Salvar o template
            template_path = os.path.join(TEMPLATES_DIR, f"{tipo_id}.docx")
            doc.save(template_path)
            print(f"Template criado: {template_path}")
    
    @staticmethod
    def gerar_documento_personalizado(tipo_id, cliente_id, dados_peticao, contraparte=None, autoridade=None, referencia_processo=None, cidade=None):
        """
        Gera um documento personalizado com base no template e nos dados fornecidos
        
        Args:
            tipo_id: ID do tipo de petição (ex: 'recurso_administrativo')
            cliente_id: ID do cliente
            dados_peticao: Dicionário com os dados da petição (fatos, fundamentos, pedidos)
            contraparte: Nome da contraparte (opcional)
            autoridade: Nome da autoridade (opcional)
            referencia_processo: Referência do processo (opcional)
            cidade: Cidade para o fechamento (opcional)
            
        Returns:
            Caminho do arquivo gerado
        """
        # Verificar se o tipo de petição existe
        if tipo_id not in TemplateManager.TIPOS_PETICAO:
            raise ValueError(f"Tipo de petição não encontrado: {tipo_id}")
        
        # Obter informações do tipo de petição
        info_tipo = TemplateManager.TIPOS_PETICAO[tipo_id]
        
        # Obter cliente
        cliente = Cliente.obter_cliente_por_id(cliente_id)
        if not cliente:
            raise ValueError(f"Cliente não encontrado: {cliente_id}")
        
        # Verificar se o template existe
        template_path = os.path.join(TEMPLATES_DIR, f"{tipo_id}.docx")
        if not os.path.exists(template_path):
            # Criar templates padrão se não existirem
            TemplateManager.criar_templates_padrao()
            
            if not os.path.exists(template_path):
                raise ValueError(f"Template não encontrado: {tipo_id}")
        
        # Carregar o template
        doc = Document(template_path)
        
        # Substituir placeholders no cabeçalho
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if "[LOGO_CLIENTE]" in paragraph.text and cliente.logo_path:
                    # Substituir texto pelo logo
                    paragraph.text = ""
                    run = paragraph.add_run()
                    if os.path.exists(cliente.logo_path):
                        run.add_picture(cliente.logo_path, width=Cm(5))
        
        # Substituir placeholders no documento
        for paragraph in doc.paragraphs:
            if "[AUTORIDADE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[AUTORIDADE]", autoridade or "PREGOEIRO")
                
            if "[REFERENCIA_PROCESSO]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[REFERENCIA_PROCESSO]", 
                                                       referencia_processo or "PROCESSO ADMINISTRATIVO Nº [NÚMERO]")
                
            if "[NOME_CLIENTE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[NOME_CLIENTE]", cliente.nome)
                
            if "[QUALIFICACAO_CLIENTE]" in paragraph.text:
                # Construir qualificação com base nos dados disponíveis
                qualificacao = ""
                
                # Usar razão social se disponível, senão usar nome
                nome_empresa = cliente.razao_social or cliente.nome
                qualificacao += nome_empresa
                
                # Adicionar CNPJ se disponível
                if cliente.cnpj and cliente.cnpj.strip() != "":
                    qualificacao += f", inscrita no CNPJ sob o nº {cliente.cnpj}"
                
                # Adicionar endereço se disponível
                if cliente.endereco_completo and cliente.endereco_completo.strip() != "":
                    qualificacao += f", com sede em {cliente.endereco_completo}"
                elif cliente.endereco and cliente.endereco.strip() != "":
                    qualificacao += f", com sede em {cliente.endereco}"
                
                # Adicionar representante legal se disponível
                if cliente.representante_legal and cliente.representante_legal.strip() != "":
                    qualificacao += f", neste ato representada por seu representante legal, {cliente.representante_legal}"
                
                # Adicionar papel do cliente no processo
                qualificacao += f", na qualidade de {info_tipo.get('papel_cliente', 'licitante')}"
                
                paragraph.text = paragraph.text.replace("[QUALIFICACAO_CLIENTE]", qualificacao)
                
            if "[CONTRAPARTE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[CONTRAPARTE]", contraparte or "[NOME DA CONTRAPARTE]")
                
            if "[FATOS]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[FATOS]", dados_peticao.get('fatos', ''))
                
            if "[FUNDAMENTOS]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[FUNDAMENTOS]", dados_peticao.get('fundamentos', ''))
                
            if "[PEDIDOS]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[PEDIDOS]", dados_peticao.get('pedidos', ''))
                
            if "[CIDADE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[CIDADE]", cidade or "São Paulo")
                
            if "[DATA]" in paragraph.text:
                data_atual = datetime.now().strftime('%d/%m/%Y')
                paragraph.text = paragraph.text.replace("[DATA]", data_atual)
                
            if "[ADVOGADO]" in paragraph.text:
                advogado = cliente.advogados[0] if cliente.advogados else "ADVOGADO"
                paragraph.text = paragraph.text.replace("[ADVOGADO]", advogado.get('nome', 'ADVOGADO'))
                
            if "[NUMERO_OAB]" in paragraph.text:
                advogado = cliente.advogados[0] if cliente.advogados else {"oab": "OAB/XX 000000"}
                paragraph.text = paragraph.text.replace("[NUMERO_OAB]", advogado.get('oab', 'OAB/XX 000000'))
        
        # Definir o nome do arquivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        tipo_sanitizado = tipo_id.replace('_', '-')
        filename = f"{tipo_sanitizado}_{cliente.nome.replace(' ', '_')}_{timestamp}.docx"
        
        # Diretório para salvar as petições
        peticoes_dir = os.path.join(os.path.dirname(__file__), 'peticoes')
        os.makedirs(peticoes_dir, exist_ok=True)
        
        # Caminho completo do arquivo
        filepath = os.path.join(peticoes_dir, filename)
        
        # Salvar o documento
        doc.save(filepath)
        
        return filepath

# Criar templates padrão se o arquivo for executado diretamente
if __name__ == "__main__":
    TemplateManager.criar_templates_padrao()
    print("Templates padrão criados com sucesso!") 