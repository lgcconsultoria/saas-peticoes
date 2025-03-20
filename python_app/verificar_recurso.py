#!/usr/bin/env python
import os
import docx
import re

def verificar_template(template_path):
    """Verifica o conteúdo do template e extrai os placeholders"""
    print(f"Verificando template: {template_path}")
    
    if not os.path.exists(template_path):
        print(f"Erro: Template não encontrado: {template_path}")
        return
    
    try:
        # Abrir o documento
        doc = docx.Document(template_path)
        
        # Extrair todo o texto
        texto_completo = ""
        for i, para in enumerate(doc.paragraphs):
            texto_completo += para.text + "\n"
            print(f"Parágrafo {i+1}: {para.text}")
        
        # Procurar por placeholders
        placeholders = re.findall(r'\[(.*?)\]', texto_completo)
        placeholders.extend(re.findall(r'##(.*?)##', texto_completo))
        
        # Remover duplicatas
        placeholders = list(set(placeholders))
        
        print("\nPlaceholders encontrados:")
        for placeholder in sorted(placeholders):
            print(f"- {placeholder}")
        
        # Verificar o parágrafo que deve conter o nome do cliente
        for i, para in enumerate(doc.paragraphs):
            if "vem, respeitosamente" in para.text:
                print(f"\nParágrafo com 'vem, respeitosamente' (#{i+1}):")
                print(para.text)
                
                # Verificar se o parágrafo contém o placeholder correto
                if "[NOME_CLIENTE]" in para.text:
                    print("✓ Contém o placeholder [NOME_CLIENTE]")
                else:
                    print("✗ NÃO contém o placeholder [NOME_CLIENTE]")
                    
                    # Sugerir correção
                    if "Cliente" in para.text:
                        print("\nSugestão de correção:")
                        print(para.text.replace("Cliente", "[NOME_CLIENTE]"))
        
    except Exception as e:
        print(f"Erro ao verificar template: {e}")

if __name__ == "__main__":
    template_path = "templates_docx/recurso_administrativo.docx"
    verificar_template(template_path) 