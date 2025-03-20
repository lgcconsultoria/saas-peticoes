from docx import Document
import os
import sys
import json
from datetime import datetime

def verificar_template(template_path):
    """Verifica o conteúdo de um template DOCX"""
    if not os.path.exists(template_path):
        print(f"Arquivo não encontrado: {template_path}")
        return
    
    doc = Document(template_path)
    print(f"Analisando template: {template_path}")
    print(f"Total de parágrafos: {len(doc.paragraphs)}")
    
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"{i}: {paragraph.text}")
    
    # Verificar placeholders
    placeholders = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text
        # Procurar por placeholders no formato [PLACEHOLDER]
        start_idx = 0
        while True:
            start_idx = text.find('[', start_idx)
            if start_idx == -1:
                break
            end_idx = text.find(']', start_idx)
            if end_idx == -1:
                break
            placeholder = text[start_idx:end_idx+1]
            placeholders.add(placeholder)
            start_idx = end_idx + 1
    
    print("\nPlaceholders encontrados:")
    for placeholder in sorted(placeholders):
        print(f"  {placeholder}")

def gerar_html_preview(tipo, fatos, argumentos, pedidos, cliente_id=None, contraparte=None, autoridade=None, referencia_processo=None, cidade=None):
    """Gera o HTML formatado para a prévia da petição"""
    try:
        # Buscar informações do cliente se cliente_id for fornecido
        cliente_nome = "Cliente"
        cliente_qualificacao = "devidamente qualificado(a) nos autos do processo em epígrafe"
        advogado_nome = "ADVOGADO"
        advogado_oab = "OAB/XX 12345"
        
        if cliente_id:
            try:
                # Carregar informações do cliente
                clientes_json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'clientes', 'clientes.json')
                if os.path.exists(clientes_json_path):
                    with open(clientes_json_path, 'r', encoding='utf-8') as f:
                        clientes = json.load(f)
                        cliente = next((c for c in clientes if str(c.get('id')) == str(cliente_id)), None)
                        if cliente:
                            cliente_nome = cliente.get('nome', 'Cliente')
                            cliente_cnpj = cliente.get('cnpj', '')
                            cliente_endereco = cliente.get('endereco', '')
                            
                            # Montar qualificação completa do cliente
                            cliente_qualificacao = f"inscrita no CNPJ sob o nº {cliente_cnpj}, com sede na {cliente_endereco}"
                            
                            # Verificar se há advogados cadastrados para o cliente
                            advogados = cliente.get('advogados', [])
                            if advogados and len(advogados) > 0:
                                advogado = advogados[0]
                                advogado_nome = advogado.get('nome', 'ADVOGADO')
                                advogado_oab = advogado.get('oab', 'OAB/XX 12345')
                            else:
                                # Usar valores padrão para advogado
                                advogado_nome = "DOUGLAS SENTURIÃO"
                                advogado_oab = "OAB/SC 73764"
            except Exception as e:
                print(f"Erro ao buscar informações do cliente: {e}")
        
        # Formatar a data atual
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        # Gerar o HTML formatado
        html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                h1, h2, h3 {{ color: #333; }}
                .text-center {{ text-align: center; }}
                .text-justify {{ text-align: justify; }}
                .text-right {{ text-align: right; }}
                .mt-4 {{ margin-top: 1.5rem; }}
                .mb-4 {{ margin-bottom: 1.5rem; }}
            </style>
        </head>
        <body>
            <h1 class="text-center">EXCELENTÍSSIMO(A) SENHOR(A) {autoridade or "JUIZ(A) DE DIREITO"}</h1>
            
            <p class="text-center mb-4"><strong>{referencia_processo or ""}</strong></p>
            
            <p class="text-justify">
                {cliente_nome}, {cliente_qualificacao}, vem, respeitosamente, à presença de Vossa Excelência, por intermédio de seu(sua) advogado(a) que esta subscreve, com fundamento na legislação aplicável, apresentar
            </p>
            
            <h2 class="text-center">{tipo.upper()}</h2>
            
            <p class="text-justify">
                em face de {contraparte or "CONTRAPARTE"}, pelos fatos e fundamentos a seguir expostos.
            </p>
            
            <h3>I - DOS FATOS</h3>
            
            <p class="text-justify">
                {fatos}
            </p>
            
            <h3>II - DOS FUNDAMENTOS</h3>
            
            <div class="text-justify">
                {argumentos}
            </div>
            
            <h3>III - DOS PEDIDOS</h3>
            
            <p class="text-justify">
                Ante o exposto, requer:
            </p>
            
            <div class="text-justify">
                {pedidos}
            </div>
            
            <p class="text-justify mt-4">
                Nestes termos,<br>
                Pede deferimento.
            </p>
            
            <p class="text-right mt-4">
                {cidade or "São Paulo"}, {data_atual}.
            </p>
            
            <p class="text-center mt-4">
                <strong>{advogado_nome}</strong><br>
                {advogado_oab}
            </p>
        </body>
        </html>
        """
        
        return html
    except Exception as e:
        print(f"Erro ao gerar HTML preview: {e}")
        return f"<html><body><h1>Erro ao gerar preview</h1><p>{str(e)}</p></body></html>"

def comparar_template_e_html():
    """Compara o template DOCX com o HTML gerado para preview"""
    print("\n=== Comparando template DOCX com HTML gerado para preview ===")
    
    # Valores de teste
    tipo = "RECURSO ADMINISTRATIVO"
    fatos = "Fatos de teste para comparação."
    argumentos = "Argumentos de teste para comparação."
    pedidos = "Pedidos de teste para comparação."
    
    # Gerar HTML
    html = gerar_html_preview(tipo, fatos, argumentos, pedidos)
    
    # Salvar HTML para inspeção
    with open("preview_test.html", "w", encoding="utf-8") as f:
        f.write(html)
    
    print(f"HTML gerado e salvo em preview_test.html")
    print("Verifique se o HTML gerado corresponde à estrutura do template DOCX.")

if __name__ == "__main__":
    template_path = os.path.join("templates_docx", "recurso_administrativo.docx")
    verificar_template(template_path)
    comparar_template_e_html() 