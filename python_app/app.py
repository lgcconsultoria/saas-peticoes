from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import os
import json
import uuid
from datetime import datetime
import requests
from docx import Document
import re
from dotenv import load_dotenv
import time
import importlib
from docx.shared import Cm
from utils.formatacao_juridica import formatar_texto_juridico, formatar_citacoes_legais
from utils.docx_generator import DocxGenerator
from utils.ai_generator import AIGenerator
from utils.validacao_juridica import ValidacaoJuridica
from openai import OpenAI
import logging

# Desativar configurações de proxy que podem estar causando problemas
os.environ.pop('HTTP_PROXY', None)
os.environ.pop('HTTPS_PROXY', None)
os.environ.pop('http_proxy', None)
os.environ.pop('https_proxy', None)

# Carregar variáveis de ambiente
load_dotenv()

app = Flask(__name__)

# Configurar CORS para permitir requisições de qualquer origem
CORS(app)

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuração adicional para CORS
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

# Configuração
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ASSISTANT_ID = os.getenv("ASSISTANT_ID")

# Log de inicialização com informações do assistente
print("="*80)
print(f"Inicializando aplicação com API OpenAI")
print(f"ID do Assistente: {ASSISTANT_ID}")

# Diretórios
PETICOES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'peticoes')
TEMPLATES_DOCX_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates_docx')
CLIENTES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'clientes')
os.makedirs(PETICOES_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DOCX_DIR, exist_ok=True)
os.makedirs(CLIENTES_DIR, exist_ok=True)

print("="*80)

# Importar OpenAI de forma segura
try:
    from openai import OpenAI
    print("Biblioteca OpenAI importada com sucesso")
except ImportError as e:
    print(f"Erro ao importar OpenAI: {e}")
    raise e

# Verificar e criar templates padrão se necessário
def verificar_templates_padrao():
    """Verifica se os templates padrão existem e os cria se necessário"""
    templates_padrao = {
        "recurso_administrativo.docx": "Recurso Administrativo",
        "impugnacao_edital.docx": "Impugnação ao Edital",
        "mandado_seguranca.docx": "Mandado de Segurança",
        "contrarrazoes_recurso.docx": "Contrarrazões de Recurso"
    }
    
    templates_existentes = [f for f in os.listdir(TEMPLATES_DOCX_DIR) 
                           if f.endswith('.docx') and not f.startswith('~$')]
    
    # Verificar se precisamos criar ou atualizar templates
    templates_para_criar = []
    for template_file in templates_padrao:
        if template_file not in templates_existentes:
            templates_para_criar.append(template_file)
    
    if templates_para_criar:
        print(f"Templates a serem criados: {', '.join(templates_para_criar)}")
        for template_file in templates_para_criar:
            tipo = templates_padrao[template_file]
            criar_template_padrao(template_file, tipo)
    else:
        print(f"Templates encontrados: {', '.join(templates_existentes)}")
        
        # Verificar se os templates existentes precisam ser atualizados
        for template_file in templates_existentes:
            if template_file in templates_padrao:
                verificar_e_atualizar_template(template_file, templates_padrao[template_file])

def criar_template_padrao(template_file, tipo):
    """Cria um template padrão para o tipo de petição especificado"""
    try:
        print(f"Criando template padrão para {tipo}...")
        
        # Criar um documento DOCX básico
        doc = Document()
        
        # Adicionar título
        p = doc.add_paragraph()
        p.alignment = 1  # Centralizado
        run = p.add_run(f"EXCELENTÍSSIMO(A) SENHOR(A) [AUTORIDADE]")
        run.bold = True
        
        # Adicionar referência do processo
        p = doc.add_paragraph()
        p.alignment = 1  # Centralizado
        run = p.add_run(f"[REFERENCIA_PROCESSO]")
        run.bold = True
        
        # Adicionar espaço para logo
        p = doc.add_paragraph()
        p.alignment = 1  # Centralizado
        p.add_run("[LOGO_CLIENTE]")
        
        # Adicionar introdução
        doc.add_paragraph("[NOME_CLIENTE], [QUALIFICACAO_CLIENTE], vem, respeitosamente, à presença de Vossa Excelência, por intermédio de seu(sua) advogado(a) que esta subscreve, apresentar")
        
        # Adicionar tipo de petição
        p = doc.add_paragraph()
        p.alignment = 1  # Centralizado
        run = p.add_run(tipo.upper())
        run.bold = True
        
        # Texto introdutório específico para cada tipo de petição
        if tipo.lower() == "recurso administrativo":
            doc.add_paragraph("pelos fatos e fundamentos a seguir expostos.")
        else:
            doc.add_paragraph("em face de [CONTRAPARTE], pelos fatos e fundamentos a seguir expostos.")
        
        # Adicionar seções
        p = doc.add_paragraph()
        run = p.add_run("I - DOS FATOS")
        run.bold = True
        
        doc.add_paragraph("[FATOS]")
        
        p = doc.add_paragraph()
        run = p.add_run("II - DOS FUNDAMENTOS")
        run.bold = True
        
        doc.add_paragraph("[FUNDAMENTOS]")
        
        p = doc.add_paragraph()
        run = p.add_run("III - DOS PEDIDOS")
        run.bold = True
        
        doc.add_paragraph("Ante o exposto, requer:")
        doc.add_paragraph("[PEDIDOS]")
        
        # Adicionar fechamento
        doc.add_paragraph("Nestes termos,")
        doc.add_paragraph("Pede deferimento.")
        doc.add_paragraph("[CIDADE], [DATA].")
        doc.add_paragraph("[ADVOGADO]")
        doc.add_paragraph("[NUMERO_OAB]")
        
        # Salvar o documento
        template_path = os.path.join(TEMPLATES_DOCX_DIR, template_file)
        doc.save(template_path)
        print(f"Template padrão criado: {template_path}")
        
    except Exception as e:
        print(f"Erro ao criar template padrão {template_file}: {e}")

def verificar_e_atualizar_template(template_file, tipo):
    """Verifica se um template existente contém todos os placeholders necessários e o atualiza se necessário"""
    try:
        template_path = os.path.join(TEMPLATES_DOCX_DIR, template_file)
        doc = Document(template_path)
        
        # Lista de placeholders que devem estar presentes
        placeholders = [
            "[FATOS]", "[FUNDAMENTOS]", "[PEDIDOS]", "[DATA]", "[CIDADE]",
            "[CONTRAPARTE]", "[AUTORIDADE]", "[REFERENCIA_PROCESSO]",
            "[LOGO_CLIENTE]", "[NOME_CLIENTE]", "[QUALIFICACAO_CLIENTE]",
            "[ADVOGADO]", "[NUMERO_OAB]"
        ]
        
        # Verificar se todos os placeholders estão presentes
        texto_completo = " ".join([p.text for p in doc.paragraphs])
        placeholders_ausentes = [p for p in placeholders if p not in texto_completo]
        
        if placeholders_ausentes:
            print(f"Template {template_file} não contém todos os placeholders necessários. Placeholders ausentes: {', '.join(placeholders_ausentes)}")
            print(f"Criando novo template para substituir o existente...")
            
            # Fazer backup do template existente
            backup_path = os.path.join(TEMPLATES_DOCX_DIR, f"backup_{template_file}")
            import shutil
            shutil.copy2(template_path, backup_path)
            print(f"Backup do template existente criado: {backup_path}")
            
            # Criar novo template
            criar_template_padrao(template_file, tipo)
        else:
            print(f"Template {template_file} contém todos os placeholders necessários.")
            
    except Exception as e:
        print(f"Erro ao verificar template {template_file}: {e}")

# Verificar e criar templates padrão
verificar_templates_padrao()

# Função para criar cliente OpenAI de forma segura
def create_openai_client():
    try:
        print("Criando cliente OpenAI com chave API...")

        os.environ.pop('HTTP_PROXY', None)
        os.environ.pop('HTTPS_PROXY', None)
        os.environ.pop('http_proxy', None)
        os.environ.pop('https_proxy', None)
        
        # Criar cliente com apenas a chave API
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        # Testar a conexão com a API
        print("Testando conexão com a API OpenAI...")
        models = client.models.list()
        print(f"Conexão bem-sucedida! {len(models.data)} modelos disponíveis.")
        
        return client
    except Exception as e:
        print(f"Erro ao criar cliente OpenAI: {e}")
        import traceback
        traceback.print_exc()
        raise e

def validar_campos_peticao(data):
    """Valida os campos obrigatórios da petição"""
    required_fields = ['tipo', 'motivo', 'fatos']
    for field in required_fields:
        if field not in data or not data[field]:
            return False, f"Campo '{field}' é obrigatório"
    return True, None

def limpar_texto(texto):
    """Remove referências e formatações indesejadas do texto"""
    if not texto:
        return ""
    
    # Remover referências do tipo 【4:6†source】 e variações
    texto_limpo = re.sub(r'【\d+:\d+†[^】]*】', '', texto)
    texto_limpo = re.sub(r'【[^】]*】', '', texto_limpo)  # Captura outras variações
    
    # Remover tags HTML
    texto_limpo = re.sub(r'<[^>]*>', '', texto_limpo)
    
    # Remover espaços extras
    texto_limpo = re.sub(r'\s+', ' ', texto_limpo).strip()
    
    return texto_limpo

def gerar_html_preview(tipo, fatos, argumentos, pedidos, cliente_id=None, cliente_nome=None, cliente_razao_social=None, cliente_cnpj=None, autoridade=None, referencia_processo=None, cidade=None):
    """Gera o HTML formatado para a prévia da petição"""
    try:
        # Buscar informações do cliente se cliente_id for fornecido
        cliente_qualificacao = ""
        advogado_nome = "ADVOGADO"
        advogado_oab = "OAB/XX 12345"
        
        # Usar cliente_nome se fornecido, senão usar valor padrão
        if not cliente_nome:
            cliente_nome = "Cliente"
        
        # Construir qualificação do cliente com os dados fornecidos
        if cliente_cnpj:
            cliente_qualificacao = f"pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {cliente_cnpj}"
        
        # Tentar obter mais informações do cliente do banco de dados
        if cliente_id:
            try:
                # Carregar informações do cliente
                clientes_json_path = os.path.join(CLIENTES_DIR, 'clientes.json')
                if os.path.exists(clientes_json_path):
                    with open(clientes_json_path, 'r', encoding='utf-8') as f:
                        clientes = json.load(f)
                        
                        # Converter cliente_id para string para comparação
                        cliente_id_str = str(cliente_id)
                        cliente = next((c for c in clientes if str(c.get('id')) == cliente_id_str), None)
                        
                        if cliente:
                            # Usar nome do cliente do banco de dados se não fornecido diretamente
                            if not cliente_nome:
                                cliente_nome = cliente.get('nome', 'Cliente')
                            
                            # Construir qualificação mais completa se tiver endereço
                            endereco = cliente.get('endereco', '')
                            if endereco and cliente_cnpj:
                                cliente_qualificacao = f"pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {cliente_cnpj}, com sede na {endereco}"
                            elif endereco:
                                cliente_qualificacao = f"pessoa jurídica de direito privado, com sede na {endereco}"
                            
                            # Obter informações do advogado
                            advogados = cliente.get('advogados', [])
                            if advogados and len(advogados) > 0:
                                advogado = advogados[0]
                                advogado_nome = advogado.get('nome', 'ADVOGADO')
                                advogado_oab = advogado.get('oab', 'OAB/XX 12345')
                        else:
                            print(f"Cliente com ID {cliente_id} não encontrado no arquivo clientes.json")
            except Exception as e:
                print(f"Erro ao carregar dados do cliente: {str(e)}")
        
        # Formatar a data atual
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        # Mapear tipos de petição para nomes de templates
        tipo_template_map = {
            "recurso": "RECURSO ADMINISTRATIVO",
            "recurso administrativo": "RECURSO ADMINISTRATIVO",
            "impugnação": "IMPUGNAÇÃO AO EDITAL",
            "impugnação ao edital": "IMPUGNAÇÃO AO EDITAL",
            "mandado": "MANDADO DE SEGURANÇA",
            "mandado de segurança": "MANDADO DE SEGURANÇA",
            "contrarrazões": "CONTRARRAZÕES DE RECURSO",
            "contrarrazões de recurso": "CONTRARRAZÕES DE RECURSO"
        }
        
        # Obter o nome do template baseado no tipo (convertido para minúsculas)
        tipo_lower = tipo.lower()
        tipo_display = tipo_template_map.get(tipo_lower, tipo.upper())
        
        # Gerar o HTML formatado seguindo a estrutura do template DOCX
        html = f"""
<div style="font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; padding: 20px; max-width: 800px;">
    <div style="text-align: center; margin-bottom: 20px;">
        <h2>EXCELENTÍSSIMO(A) SENHOR(A) {autoridade}</h2>
    </div>
    
    <div style="text-align: center; margin-bottom: 20px;">
        <strong>{referencia_processo or ""}</strong>
    </div>
    
    <div style="text-align: justify; margin-bottom: 20px;">
        {cliente_nome}, {cliente_qualificacao}, vem, respeitosamente, à presença de Vossa Excelência, por intermédio de representante legal que esta subscreve, apresentar
    </div>
    
    <div style="text-align: center; margin-bottom: 20px;">
        <h3>{tipo_display}</h3>
    </div>
    
    <div style="text-align: justify; margin-bottom: 20px;">
        em face da decisão proferida por esta autoridade, pelos fatos e fundamentos a seguir expostos.
    </div>
    
    <div style="margin-bottom: 20px;">
        <h3>I - DOS FATOS</h3>
        <div style="text-align: justify;">
            {fatos}
        </div>
    </div>
    
    <div style="margin-bottom: 20px;">
        <h3>II - DOS FUNDAMENTOS</h3>
        <div style="text-align: justify;">
            {argumentos}
        </div>
    </div>
    
    <div style="margin-bottom: 20px;">
        <h3>III - DOS PEDIDOS</h3>
        <div style="text-align: justify;">
            Ante o exposto, requer:
        </div>
        <div style="text-align: justify; margin-left: 20px;">
            {pedidos}
        </div>
    </div>
    
    <div style="text-align: justify; margin-top: 30px;">
        Nestes termos,<br>
        Pede deferimento.
    </div>
    
    <div style="text-align: right; margin-top: 30px;">
        {cidade or "São Paulo"}, {data_atual}.
    </div>
    
    <div style="text-align: center; margin-top: 50px;">
        <strong>{advogado_nome}</strong><br>
        {advogado_oab}
    </div>
</div>
"""
        
        return html
    except Exception as e:
        print(f"Erro ao gerar HTML preview: {e}")
        # Em caso de erro, retornar um HTML básico
        return f"""
<div style="font-family: Arial, sans-serif; padding: 20px;">
    <h2 style="color: red;">Erro ao gerar prévia</h2>
    <p>Ocorreu um erro ao gerar a prévia da petição: {str(e)}</p>
</div>
"""

# Inicializar geradores
docx_generator = DocxGenerator(TEMPLATES_DOCX_DIR, PETICOES_DIR, CLIENTES_DIR)
ai_generator = AIGenerator(api_key=os.getenv("OPENAI_API_KEY"), assistant_id=os.getenv("ASSISTANT_ID"))

# Inicializar validador jurídico
validador_juridico = ValidacaoJuridica()

def gerar_peticao_com_assistente(tipo, motivo, fatos):
    """Gera uma petição usando o Assistente da OpenAI"""
    try:
        print("="*80)
        print(f"Iniciando geração de petição com o Assistente da OpenAI...")
        
        # Usar a nova classe AIGenerator
        resultado = ai_generator.gerar_peticao(tipo, motivo, fatos)
        
        return resultado
    except Exception as error:
        print(f"Erro ao gerar petição com assistente: {error}")
        raise error

def gerar_docx(tipo, texto, cliente_id=None, cliente_nome=None, cliente_razao_social=None, cliente_cnpj=None, autoridade=None, referencia_processo=None, cidade=None, fatos_texto=None, argumentos_texto=None, pedidos_texto=None):
    """Gera um documento DOCX com o texto da petição"""
    try:
        print(f"Iniciando geração do documento DOCX para petição: {tipo}")
        
        # Preparar dados para o gerador de documentos
        dados_peticao = {
            'cliente_id': cliente_id,
            'cliente_nome': cliente_nome,
            'cliente_razao_social': cliente_razao_social,
            'cliente_cnpj': cliente_cnpj,
            'autoridade': autoridade,
            'referencia_processo': referencia_processo,
            'cidade': cidade
        }
        
        # Verificar se os textos foram fornecidos diretamente
        if fatos_texto is not None or argumentos_texto is not None or pedidos_texto is not None:
            fatos = fatos_texto or ""
            fundamentos = argumentos_texto or ""
            pedidos = pedidos_texto or ""
        else:
            # Se texto não estiver vazio, extrair fatos, fundamentos e pedidos
            if texto:
                # Extrair fatos, fundamentos e pedidos do texto
                fatos_match = re.search(r'I - DOS FATOS\s*([\s\S]*?)(?=II - DOS FUNDAMENTOS|$)', texto)
                fundamentos_match = re.search(r'II - DOS FUNDAMENTOS\s*([\s\S]*?)(?=III - DOS PEDIDOS|$)', texto)
                pedidos_match = re.search(r'III - DOS PEDIDOS\s*(?:Ante o exposto, requer:)?\s*([\s\S]*?)(?=Nestes termos|$)', texto)
                
                fatos = fatos_match.group(1).strip() if fatos_match else ""
                fundamentos = fundamentos_match.group(1).strip() if fundamentos_match else ""
                pedidos = pedidos_match.group(1).strip() if pedidos_match else ""
                
                # Limpar os textos
                fatos = limpar_texto(fatos)
                fundamentos = limpar_texto(fundamentos)
                pedidos = limpar_texto(pedidos)
            else:
                fatos = ""
                fundamentos = ""
                pedidos = ""
        
        # Adicionar textos aos dados da petição
        dados_peticao['fatos_texto'] = fatos
        dados_peticao['argumentos_texto'] = fundamentos
        dados_peticao['pedidos_texto'] = pedidos
        
        # Gerar documento usando o DocxGenerator
        filepath = docx_generator.gerar_documento(tipo, dados_peticao)
        
        return filepath
    except Exception as e:
        print(f"Erro ao gerar documento DOCX: {e}")
        import traceback
        traceback.print_exc()
        raise e

@app.route('/api/gerar-peticao', methods=['POST'])
def api_gerar_peticao():
    """Endpoint para gerar petições"""
    print("="*80)
    print("Recebida solicitação para gerar petição")
    print("Headers:", dict(request.headers))
    print("Dados recebidos:", request.get_json())
    
    try:
        data = request.json
        
        # Verificar se o cliente foi enviado no formato antigo (objeto cliente) e extrair o cliente_id
        if 'cliente' in data and isinstance(data['cliente'], dict) and 'id' in data['cliente']:
            cliente_obj = data['cliente']
            data['cliente_id'] = cliente_obj.get('id')
            data['cliente_nome'] = cliente_obj.get('nome')
            data['cliente_razao_social'] = cliente_obj.get('razaoSocial')
            data['cliente_cnpj'] = cliente_obj.get('cnpj')
            print(f"Cliente extraído do formato antigo: {data['cliente_id']}")
        
        # Garantir que temos as informações do cliente
        if data.get('cliente_id'):
            try:
                # Carregar informações do cliente
                clientes_json_path = os.path.join(CLIENTES_DIR, 'clientes.json')
                if os.path.exists(clientes_json_path):
                    with open(clientes_json_path, 'r', encoding='utf-8') as f:
                        clientes = json.load(f)
                        
                        # Converter cliente_id para string para comparação
                        cliente_id_str = str(data.get('cliente_id'))
                        cliente = next((c for c in clientes if str(c.get('id')) == cliente_id_str), None)
                        
                        if cliente:
                            # Usar o nome do cliente do banco de dados se não fornecido diretamente
                            if not data.get('cliente_nome'):
                                data['cliente_nome'] = cliente.get('nome')
            except Exception as e:
                print(f"Erro ao carregar dados do cliente: {str(e)}")
        
        # Validar campos
        valid, error_msg = validar_campos_peticao(data)
        if not valid:
            print(f"Erro de validação: {error_msg}")
            return jsonify({
                "error": "Campos inválidos",
                "message": error_msg
            }), 400
        
        # Verificar se temos um ASSISTANT_ID configurado
        if not ASSISTANT_ID:
            print("Erro: ID do assistente não configurado")
            return jsonify({
                "error": "ID do assistente não fornecido",
                "message": "É necessário fornecer o ID do assistente para gerar a petição."
            }), 400
            
        print(f"Usando ID do assistente: {ASSISTANT_ID}")
        
        # Obter o tipo de petição e normalizar
        tipo_original = data.get('tipo')
        motivo = data.get('motivo')
        fatos = data.get('fatos')
        
        # Mapear tipos abreviados para nomes completos
        tipo_map = {
            "recurso": "Recurso Administrativo",
            "impugnação": "Impugnação ao Edital",
            "mandado": "Mandado de Segurança",
            "contrarrazões": "Contrarrazões de Recurso"
        }
        
        # Normalizar o tipo para garantir consistência
        tipo = tipo_map.get(tipo_original.lower(), tipo_original)
        print(f"Tipo de petição normalizado: {tipo_original} -> {tipo}")
        
        # Gerar petição com o assistente
        print(f"Iniciando geração de petição: {tipo}")
        print(f"Motivo fornecido: {motivo}")
        print(f"Fatos básicos fornecidos: {fatos}")
        
        # Adicionar contexto adicional se disponível
        contexto_adicional = None
        if data.get('processo'):
            contexto_adicional = f"Número do processo: {data.get('processo')}. "
        if data.get('orgao'):
            contexto_adicional = (contexto_adicional or "") + f"Órgão/Entidade: {data.get('orgao')}. "
        if data.get('autoridade'):
            contexto_adicional = (contexto_adicional or "") + f"Autoridade: {data.get('autoridade')}. "
        
        # Gerar petição com conteúdo expandido
        resultado = gerar_peticao_com_assistente(tipo, motivo, fatos)
        print(f"Petição gerada com sucesso: {tipo}")
        
        # Extrair as partes da petição
        fatos_texto = resultado.get("fatos", "")
        argumentos_texto = resultado.get("argumentos", "")
        pedido_texto = resultado.get("pedido", "")
        
        # Verificar se o conteúdo foi gerado corretamente
        if not fatos_texto or len(fatos_texto.strip()) < 50:
            print(f"AVISO: Fatos gerados muito curtos ou vazios: '{fatos_texto}'")
            # Tentar gerar novamente com mais ênfase nos fatos
            try:
                print("Tentando gerar novamente com mais ênfase nos fatos...")
                resultado = ai_generator.gerar_peticao(tipo, motivo, fatos, contexto_adicional=f"É EXTREMAMENTE IMPORTANTE expandir os fatos fornecidos em uma narrativa jurídica completa. Fatos básicos: {fatos}")
                fatos_texto = resultado.get("fatos", fatos_texto)
                argumentos_texto = resultado.get("argumentos", argumentos_texto)
                pedido_texto = resultado.get("pedido", pedido_texto)
            except Exception as e:
                print(f"Erro ao tentar gerar novamente: {e}")
        
        if not argumentos_texto or len(argumentos_texto.strip()) < 100:
            print(f"AVISO: Argumentos gerados muito curtos ou vazios: '{argumentos_texto}'")
            # Tentar gerar novamente com mais ênfase nos argumentos
            try:
                print("Tentando gerar novamente com mais ênfase nos argumentos...")
                resultado_args = ai_generator.gerar_peticao(tipo, motivo, fatos, contexto_adicional=f"É EXTREMAMENTE IMPORTANTE fornecer argumentos jurídicos detalhados com citações de leis e jurisprudências.")
                argumentos_texto = resultado_args.get("argumentos", argumentos_texto)
                if not fatos_texto:
                    fatos_texto = resultado_args.get("fatos", "")
                if not pedido_texto:
                    pedido_texto = resultado_args.get("pedido", "")
            except Exception as e:
                print(f"Erro ao tentar gerar novamente: {e}")
        
        # Se alguma parte ainda estiver vazia, retornar erro
        if not fatos_texto or not argumentos_texto or not pedido_texto:
            print("Erro: Conteúdo da petição incompleto após múltiplas tentativas")
            return jsonify({
                "error": "Conteúdo incompleto",
                "message": "Não foi possível gerar o conteúdo completo da petição após múltiplas tentativas."
            }), 500
        
        # Limpar os textos
        fatos_texto = limpar_texto(fatos_texto)
        argumentos_texto = limpar_texto(argumentos_texto)
        pedido_texto = limpar_texto(pedido_texto)
        
        # Gerar HTML para preview
        html_preview = gerar_html_preview(
            tipo=tipo,
            fatos=fatos_texto,
            argumentos=argumentos_texto,
            pedidos=pedido_texto,
            cliente_id=data.get('cliente_id'),
            cliente_nome=data.get('cliente_nome'),
            cliente_razao_social=data.get('cliente_razao_social'),
            cliente_cnpj=data.get('cliente_cnpj'),
            autoridade=data.get('autoridade'),
            referencia_processo=data.get('processo'),
            cidade=None  # Cidade será obtida do cliente ou padrão
        )
        
        # Gerar documento DOCX
        docx_filename = gerar_docx(
            tipo=tipo,
            texto=None,  # Não usar o texto completo, mas as partes separadas
            cliente_id=data.get('cliente_id'),
            cliente_nome=data.get('cliente_nome'),
            cliente_razao_social=data.get('cliente_razao_social'),
            cliente_cnpj=data.get('cliente_cnpj'),
            autoridade=data.get('autoridade'),
            referencia_processo=data.get('processo'),
            cidade=None,  # Cidade será obtida do cliente ou padrão
            fatos_texto=fatos_texto,
            argumentos_texto=argumentos_texto,
            pedidos_texto=pedido_texto
        )
        
        # Construir URL para download
        download_url = f"/download/{docx_filename}"
        
        # Retornar resposta
        return jsonify({
            "success": True,
            "message": "Petição gerada com sucesso",
            "preview": html_preview,
            "download_url": download_url,
            "tipo": tipo,
            "fatos": fatos_texto,
            "argumentos": argumentos_texto,
            "pedidos": pedido_texto
        })
        
    except Exception as error:
        print(f"Erro ao gerar petição: {error}")
        import traceback
        traceback.print_exc()
        
        return jsonify({
            "error": "Erro ao gerar petição",
            "message": str(error)
        }), 500

@app.route('/api/peticoes', methods=['GET'])
def api_listar_peticoes():
    """Endpoint para listar petições"""
    try:
        if not os.path.exists(PETICOES_DIR):
            return jsonify({
                "sucesso": True,
                "peticoes": []
            })
        
        arquivos = os.listdir(PETICOES_DIR)
        peticoes = [
            {
                "nome": arquivo,
                "caminho": os.path.join(PETICOES_DIR, arquivo)
            }
            for arquivo in arquivos if arquivo.endswith('.docx')
        ]
        
        return jsonify({
            "sucesso": True,
            "peticoes": peticoes
        })
    except Exception as error:
        return jsonify({
            "error": "Falha ao listar petições",
            "message": str(error)
        }), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Endpoint para download de arquivos"""
    try:
        print(f"="*80)
        print(f"Solicitação de download para arquivo: {filename}")
        
        # Sanitizar o nome do arquivo para evitar path traversal
        filename = os.path.basename(filename)
        
        # Verificar se o arquivo existe
        filepath = os.path.join(PETICOES_DIR, filename)
        print(f"Caminho completo do arquivo: {filepath}")
        
        if not os.path.exists(filepath):
            print(f"Erro: Arquivo não encontrado: {filepath}")
            return jsonify({
                "error": "Arquivo não encontrado",
                "message": "O arquivo solicitado não existe no servidor."
            }), 404
        
        if not os.path.isfile(filepath):
            print(f"Erro: O caminho não é um arquivo: {filepath}")
            return jsonify({
                "error": "Arquivo inválido",
                "message": "O caminho especificado não é um arquivo válido."
            }), 400
        
        # Verificar se o arquivo é um DOCX
        if not filename.lower().endswith('.docx'):
            print(f"Erro: Tipo de arquivo inválido: {filename}")
            return jsonify({
                "error": "Tipo de arquivo inválido",
                "message": "Apenas arquivos DOCX são permitidos."
            }), 400
        
        print(f"Arquivo encontrado, iniciando download: {filename}")
        
        # Enviar o arquivo para download com o tipo MIME correto e headers específicos
        response = send_from_directory(
            PETICOES_DIR, 
            filename, 
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Adicionar headers específicos para download
        response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
        response.headers["Access-Control-Expose-Headers"] = "Content-Disposition"
        response.headers["Access-Control-Allow-Origin"] = "*"
        
        return response
        
    except Exception as e:
        print(f"Erro ao processar download: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": "Erro ao processar download",
            "message": str(e)
        }), 500
    finally:
        print("="*80)

@app.route('/')
def index():
    """Página inicial"""
    return render_template('index.html')

@app.route('/api/status', methods=['GET'])
def check_status():
    try:
        # Verificar conexão com OpenAI
        client = OpenAI()
        models = client.models.list()
        
        return jsonify({
            'status': 'online',
            'message': 'API está funcionando normalmente',
            'openai_connected': True,
            'timestamp': datetime.now().isoformat()
        }), 200
    except Exception as e:
        logger.error(f"Erro ao verificar status: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': str(e),
            'openai_connected': False,
            'timestamp': datetime.now().isoformat()
        }), 500

@app.route('/api/clientes', methods=['GET'])
def api_listar_clientes():
    """Endpoint para listar os clientes disponíveis"""
    print("="*80)
    print("Recebida solicitação para listar clientes")
    
    try:
        from templates_manager import Cliente
        
        # Verificar se o arquivo de clientes existe
        import os
        clientes_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'clientes.json')
        if not os.path.exists(clientes_file):
            print(f"Arquivo de clientes não encontrado: {clientes_file}")
            # Tentar carregar do diretório clientes
            clientes_file_alt = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'clientes', 'clientes.json')
            if os.path.exists(clientes_file_alt):
                print(f"Usando arquivo alternativo: {clientes_file_alt}")
                # Copiar o arquivo para o diretório data
                import shutil
                os.makedirs(os.path.dirname(clientes_file), exist_ok=True)
                shutil.copy2(clientes_file_alt, clientes_file)
                print(f"Arquivo copiado para: {clientes_file}")
        
        clientes = Cliente.carregar_clientes()
        print(f"Clientes carregados: {len(clientes)}")
        
        # Converter objetos Cliente para dicionários
        clientes_data = []
        for cliente in clientes:
            cliente_dict = {
                "id": cliente.id,
                "nome": cliente.nome,
                "cnpj": cliente.cnpj,
                "endereco": cliente.endereco,
                "advogados": cliente.advogados
            }
            clientes_data.append(cliente_dict)
            print(f"Cliente adicionado: {cliente.nome}")
        
        print(f"Total de clientes retornados: {len(clientes_data)}")
        
        return jsonify({
            "success": True,
            "clientes": clientes_data
        })
    except Exception as error:
        import traceback
        print(f"Erro ao listar clientes: {error}")
        print(traceback.format_exc())
        return jsonify({
            "error": "Erro ao listar clientes",
            "message": str(error)
        }), 500

@app.route('/api/tipos-peticao', methods=['GET'])
def api_listar_tipos_peticao():
    """Endpoint para listar os tipos de petição disponíveis"""
    try:
        from templates_manager import TemplateManager
        
        tipos = {}
        for tipo_id, info in TemplateManager.TIPOS_PETICAO.items():
            tipos[tipo_id] = {
                "id": tipo_id,
                "titulo": info['titulo'],
                "papel_cliente": info['papel_cliente'],
                "papel_contraparte": info['papel_contraparte']
            }
        
        return jsonify({
            "success": True,
            "tipos": tipos
        })
    except Exception as error:
        print(f"Erro ao listar tipos de petição: {error}")
        return jsonify({
            "error": "Erro ao listar tipos de petição",
            "message": str(error)
        }), 500

@app.route('/api/validar-peticao', methods=['POST'])
def api_validar_peticao():
    """Endpoint para validar petições"""
    try:
        # Obter dados da requisição
        data = request.json
        
        # Validar campos obrigatórios
        if not data or not isinstance(data, dict):
            return jsonify({
                "erro": "Dados inválidos",
                "mensagem": "Os dados da petição não foram fornecidos corretamente."
            }), 400
        
        # Validar a petição
        relatorio = validador_juridico.gerar_relatorio_validacao(data)
        
        return jsonify({
            "valido": relatorio['valido'],
            "erros": relatorio['erros'],
            "estatisticas": relatorio['estatisticas'],
            "citacoes": relatorio['citacoes_encontradas']
        })
    
    except Exception as error:
        print(f"Erro ao validar petição: {error}")
        return jsonify({
            "erro": "Erro interno",
            "mensagem": str(error)
        }), 500

if __name__ == '__main__':
    port = int(os.getenv("PORT", 5000))
    print(f"Iniciando servidor Flask na porta {port}...")
    app.run(debug=True, host='0.0.0.0', port=port) 